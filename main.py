# Legal Document Q&A Assistant with RAG
# Complete system for legal document analysis and Q&A

import os
import streamlit as st
import pandas as pd
from typing import List, Dict, Any, Optional
import PyPDF2
import docx
from io import BytesIO
import re
from datetime import datetime
import json
import hashlib

# Vector Database and Embeddings
import pinecone
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np

# LLM Integration
import openai
from transformers import AutoTokenizer, AutoModelForCausalLM
import torch

# Document Processing
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
import spacy

# Web Framework
import asyncio
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# Configuration
class Config:
    # Vector DB Settings
    PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
    PINECONE_ENVIRONMENT = os.getenv("PINECONE_ENVIRONMENT", "us-west1-gcp")
    PINECONE_INDEX_NAME = "legal-docs"
    
    # OpenAI Settings
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
    
    # Model Settings
    EMBEDDING_MODEL = "all-MiniLM-L6-v2"
    LLM_MODEL = "gpt-3.5-turbo"
    CHUNK_SIZE = 1000
    CHUNK_OVERLAP = 200
    
    # Legal-specific settings
    RISK_KEYWORDS = [
        "unlimited liability", "personal guarantee", "non-compete", 
        "confidentiality breach", "penalty", "termination", 
        "indemnification", "force majeure", "breach of contract",
        "liquidated damages", "arbitration", "governing law"
    ]

config = Config()

class DocumentProcessor:
    """Handles document ingestion and processing"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Load spaCy model for legal text processing
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except IOError:
            st.warning("spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp = None
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {str(e)}")
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error extracting DOCX text: {str(e)}")
    
    def preprocess_legal_text(self, text: str) -> str:
        """Clean and preprocess legal text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page numbers and headers/footers
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        
        # Clean up common legal document artifacts
        text = re.sub(r'_+', '', text)  # Remove underscores
        text = re.sub(r'-{2,}', '', text)  # Remove multiple dashes
        
        return text.strip()
    
    def extract_metadata(self, text: str, filename: str) -> Dict[str, Any]:
        """Extract metadata from legal document"""
        metadata = {
            "filename": filename,
            "upload_date": datetime.now().isoformat(),
            "word_count": len(text.split()),
            "char_count": len(text),
            "document_type": self.classify_document_type(text),
            "parties": self.extract_parties(text),
            "key_dates": self.extract_dates(text),
            "risk_indicators": self.identify_risk_indicators(text)
        }
        return metadata
    

    
    def extract_parties(self, text: str) -> List[str]:
        """Extract party names from legal document"""
        parties = []
        if self.nlp:
            doc = self.nlp(text[:2000])  # Process first 2000 chars for efficiency
            for ent in doc.ents:
                if ent.label_ in ["ORG", "PERSON"] and len(ent.text) > 2:
                    parties.append(ent.text)
        
        # Also look for common legal patterns
        party_patterns = [
            r'"([^"]+)"[,\s]+(?:a|an)\s+(?:corporation|company|llc|ltd)',
            r'between\s+([^,]+),?\s+and\s+([^,\n]+)',
        ]
        
        for pattern in party_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    parties.extend(match)
                else:
                    parties.append(match)
        
        return list(set(parties[:10]))  # Return unique parties, max 10
    
    def extract_dates(self, text: str) -> List[str]:
        """Extract important dates from document"""
        date_patterns = [
            r'\b\d{1,2}/\d{1,2}/\d{4}\b',
            r'\b\d{1,2}-\d{1,2}-\d{4}\b',
            r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b'
        ]
        
        dates = []
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            dates.extend(matches)
        
        return list(set(dates[:10]))  # Return unique dates, max 10
    
    def identify_risk_indicators(self, text: str) -> List[str]:
        """Identify potential risk indicators in the document"""
        text_lower = text.lower()
        found_risks = []
        
        for risk in config.RISK_KEYWORDS:
            if risk.lower() in text_lower:
                found_risks.append(risk)
        
        return found_risks
    
    def process_document(self, file_content: bytes, filename: str) -> List[Document]:
        """Process a document and return chunks with metadata"""
        # Extract text based on file type
        if filename.lower().endswith('.pdf'):
            text = self.extract_text_from_pdf(file_content)
        elif filename.lower().endswith('.docx'):
            text = self.extract_text_from_docx(file_content)
        else:
            text = file_content.decode('utf-8')
        
        # Preprocess text
        text = self.preprocess_legal_text(text)
        
        # Extract metadata
        metadata = self.extract_metadata(text, filename)
        
        # Split into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create Document objects with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            doc_metadata = metadata.copy()
            doc_metadata.update({
                "chunk_id": i,
                "total_chunks": len(chunks),
                "chunk_hash": hashlib.md5(chunk.encode()).hexdigest()
            })
            
            documents.append(Document(
                page_content=chunk,
                metadata=doc_metadata
            ))
        
        return documents

class VectorStore:
    """Handles vector storage and retrieval using multiple backends"""
    
    def __init__(self, backend="faiss"):
        self.backend = backend
        self.embedding_model = SentenceTransformer(config.EMBEDDING_MODEL)
        
        if backend == "pinecone":
            self.setup_pinecone()
        elif backend == "faiss":
            self.setup_faiss()
        else:
            raise ValueError(f"Unsupported backend: {backend}")
        
        self.documents = []  # Store original documents for retrieval
    
    def setup_pinecone(self):
        """Initialize Pinecone vector database"""
        if not config.PINECONE_API_KEY:
            raise ValueError("PINECONE_API_KEY not found in environment")
        
        pinecone.init(
            api_key=config.PINECONE_API_KEY,
            environment=config.PINECONE_ENVIRONMENT
        )
        
        # Create index if it doesn't exist
        if config.PINECONE_INDEX_NAME not in pinecone.list_indexes():
            pinecone.create_index(
                config.PINECONE_INDEX_NAME,
                dimension=384,  # all-MiniLM-L6-v2 dimension
                metric="cosine"
            )
        
        self.index = pinecone.Index(config.PINECONE_INDEX_NAME)
    
    def setup_faiss(self):
        """Initialize FAISS vector database"""
        self.dimension = 384  # all-MiniLM-L6-v2 dimension
        self.index = faiss.IndexFlatIP(self.dimension)
        self.id_to_doc = {}
    
    def add_documents(self, documents: List[Document]):
        """Add documents to vector store"""
        texts = [doc.page_content for doc in documents]
        embeddings = self.embedding_model.encode(texts)
        
        if self.backend == "pinecone":
            # Prepare vectors for Pinecone
            vectors = []
            for i, (doc, embedding) in enumerate(zip(documents, embeddings)):
                vector_id = f"doc_{len(self.documents) + i}"
                vectors.append({
                    "id": vector_id,
                    "values": embedding.tolist(),
                    "metadata": doc.metadata
                })
            
            # Upsert to Pinecone
            self.index.upsert(vectors)
        
        elif self.backend == "faiss":
            # Add to FAISS
            embeddings = np.array(embeddings).astype('float32')
            faiss.normalize_L2(embeddings)
            
            start_id = len(self.documents)
            self.index.add(embeddings)
            
            # Store document mapping
            for i, doc in enumerate(documents):
                self.id_to_doc[start_id + i] = doc
        
        # Store documents
        self.documents.extend(documents)
    
    def search(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """Search for similar documents"""
        query_embedding = self.embedding_model.encode([query])
        
        if self.backend == "pinecone":
            results = self.index.query(
                vector=query_embedding[0].tolist(),
                top_k=top_k,
                include_metadata=True
            )
            
            search_results = []
            for match in results['matches']:
                search_results.append({
                    "content": None,  # Need to retrieve from stored documents
                    "metadata": match['metadata'],
                    "score": match['score']
                })
        
        elif self.backend == "faiss":
            query_embedding = np.array(query_embedding).astype('float32')
            faiss.normalize_L2(query_embedding)
            
            scores, indices = self.index.search(query_embedding, top_k)
            
            search_results = []
            for score, idx in zip(scores[0], indices[0]):
                if idx in self.id_to_doc:
                    doc = self.id_to_doc[idx]
                    search_results.append({
                        "content": doc.page_content,
                        "metadata": doc.metadata,
                        "score": float(score)
                    })
        
        return search_results

class LegalLLM:
    """Handles LLM interactions for legal Q&A"""
    
    def __init__(self, model_name: str = config.LLM_MODEL):
        self.model_name = model_name
        
        if model_name.startswith("gpt"):
            openai.api_key = config.OPENAI_API_KEY
            self.provider = "openai"
        else:
            # For local models
            self.provider = "local"
            self.tokenizer = AutoTokenizer.from_pretrained(model_name)
            self.model = AutoModelForCausalLM.from_pretrained(model_name)
    
    def generate_response(self, query: str, context: List[Dict[str, Any]], 
                         task_type: str = "qa") -> str:
        """Generate response based on query and context"""
        
        # Prepare context
        context_text = ""
        for i, doc in enumerate(context[:3]):  # Use top 3 results
            context_text += f"Document {i+1} ({doc['metadata']['filename']}):\n"
            context_text += f"{doc['content'][:500]}...\n\n"
        
        # Create prompt based on task type
        if task_type == "qa":
            prompt = self.create_qa_prompt(query, context_text)
        elif task_type == "summarize":
            prompt = self.create_summarize_prompt(context_text)
        elif task_type == "compare":
            prompt = self.create_compare_prompt(query, context_text)
        elif task_type == "risk_analysis":
            prompt = self.create_risk_analysis_prompt(context_text)
        else:
            prompt = self.create_qa_prompt(query, context_text)
        
        # Generate response
        if self.provider == "openai":
            response = openai.ChatCompletion.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "You are a legal assistant AI. Provide accurate, helpful information while noting that this is not legal advice and users should consult with qualified attorneys."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000,
                temperature=0.1
            )
            return response.choices[0].message.content
        
        else:
            # For local models
            inputs = self.tokenizer.encode(prompt, return_tensors="pt")
            with torch.no_grad():
                outputs = self.model.generate(
                    inputs,
                    max_length=inputs.shape[1] + 500,
                    temperature=0.1,
                    pad_token_id=self.tokenizer.eos_token_id
                )
            
            response = self.tokenizer.decode(outputs[0][inputs.shape[1]:], 
                                           skip_special_tokens=True)
            return response
    
    def create_qa_prompt(self, query: str, context: str) -> str:
        return f"""Based on the following legal document excerpts, answer the question. 
Be precise and cite specific sections when possible.

Context:
{context}

Question: {query}

Answer:"""
    
    def create_summarize_prompt(self, context: str) -> str:
        return f"""Summarize the key points of this legal document. Include:
- Main parties involved
- Key terms and conditions
- Important dates and deadlines
- Notable clauses or provisions

Document:
{context}

Summary:"""
    
    def create_compare_prompt(self, query: str, context: str) -> str:
        return f"""Compare the following legal documents focusing on: {query}

Documents:
{context}

Comparison:"""
    
    def create_risk_analysis_prompt(self, context: str) -> str:
        return f"""Analyze the following legal document for potential risks and concerns:

Document:
{context}

Risk Analysis:
1. High-risk clauses:
2. Potential liabilities:
3. Recommended actions:"""

class LegalRAGSystem:
    """Main RAG system that coordinates all components"""
    
    def __init__(self, vector_backend="faiss"):
        self.doc_processor = DocumentProcessor()
        self.vector_store = VectorStore(backend=vector_backend)
        self.llm = LegalLLM()
        self.document_index = {}  # Track uploaded documents
    
    def upload_documents(self, files: List[tuple]) -> Dict[str, Any]:
        """Upload and process multiple documents"""
        results = {
            "processed": 0,
            "failed": 0,
            "errors": [],
            "documents": []
        }
        
        for file_content, filename in files:
            try:
                # Process document
                documents = self.doc_processor.process_document(file_content, filename)
                
                # Add to vector store
                self.vector_store.add_documents(documents)
                
                # Update index
                doc_id = hashlib.md5(f"{filename}{datetime.now()}".encode()).hexdigest()
                self.document_index[doc_id] = {
                    "filename": filename,
                    "upload_date": datetime.now().isoformat(),
                    "chunk_count": len(documents),
                    "metadata": documents[0].metadata if documents else {}
                }
                
                results["processed"] += 1
                results["documents"].append({
                    "id": doc_id,
                    "filename": filename,
                    "chunks": len(documents)
                })
                
            except Exception as e:
                results["failed"] += 1
                results["errors"].append(f"{filename}: {str(e)}")
        
        return results
    
    def query(self, question: str, task_type: str = "qa") -> Dict[str, Any]:
        """Process a query and return response with sources"""
        
        # Search for relevant documents
        search_results = self.vector_store.search(question, top_k=5)
        
        if not search_results:
            return {
                "response": "I couldn't find relevant information in the uploaded documents. Please make sure you've uploaded legal documents and try rephrasing your question.",
                "sources": [],
                "confidence": 0.0
            }
        
        # Generate response using LLM
        response = self.llm.generate_response(question, search_results, task_type)
        
        # Prepare sources
        sources = []
        for result in search_results[:3]:
            sources.append({
                "filename": result["metadata"]["filename"],
                "document_type": result["metadata"].get("document_type", "Unknown"),
                "relevance_score": result["score"],
                "content_preview": result["content"][:200] + "..." if result["content"] else ""
            })
        
        return {
            "response": response,
            "sources": sources,
            "confidence": search_results[0]["score"] if search_results else 0.0,
            "task_type": task_type
        }
    
    def get_document_summary(self, doc_id: str = None) -> Dict[str, Any]:
        """Get summary of uploaded documents"""
        if doc_id and doc_id in self.document_index:
            return self.document_index[doc_id]
        
        return {
            "total_documents": len(self.document_index),
            "total_chunks": sum(doc["chunk_count"] for doc in self.document_index.values()),
            "documents": list(self.document_index.values())
        }

# Streamlit Web Interface
def create_streamlit_app():
    """Create the Streamlit web application"""
    
    st.set_page_config(
        page_title="Legal Document Q&A Assistant",
        page_icon="⚖️",
        layout="wide"
    )
    
    st.title("⚖️ Legal Document Q&A Assistant")
    st.markdown("Upload legal documents and ask questions in plain English")
    
    # Initialize session state
    if "rag_system" not in st.session_state:
        st.session_state.rag_system = LegalRAGSystem()
    
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    
    # Sidebar for document management
    with st.sidebar:
        st.header("📄 Document Management")
        
        uploaded_files = st.file_uploader(
            "Upload Legal Documents",
            accept_multiple_files=True,
            type=['pdf', 'docx', 'txt']
        )
        
        if uploaded_files:
            if st.button("Process Documents"):
                with st.spinner("Processing documents..."):
                    files = [(file.read(), file.name) for file in uploaded_files]
                    results = st.session_state.rag_system.upload_documents(files)
                
                if results["processed"] > 0:
                    st.success(f"✅ Processed {results['processed']} documents")
                
                if results["failed"] > 0:
                    st.error(f"❌ Failed to process {results['failed']} documents")
                    for error in results["errors"]:
                        st.error(error)
        
        # Document summary
        st.header("📊 Document Summary")
        summary = st.session_state.rag_system.get_document_summary()
        st.metric("Total Documents", summary["total_documents"])
        st.metric("Total Chunks", summary["total_chunks"])
        
        if summary["documents"]:
            st.subheader("Uploaded Documents:")
            for doc in summary["documents"]:
                with st.expander(doc["filename"]):
                    st.write(f"**Type:** {doc['metadata'].get('document_type', 'Unknown')}")
                    st.write(f"**Chunks:** {doc['chunk_count']}")
                    st.write(f"**Upload Date:** {doc['upload_date']}")
                    if doc['metadata'].get('risk_indicators'):
                        st.write(f"**Risk Indicators:** {', '.join(doc['metadata']['risk_indicators'])}")
    
    # Main chat interface
    col1, col2 = st.columns([3, 1])
    
    with col1:
        st.header("💬 Ask Questions")
        
        # Task type selection
        task_type = st.selectbox(
            "Select Task Type:",
            ["qa", "summarize", "compare", "risk_analysis"],
            format_func=lambda x: {
                "qa": "Q&A - Ask specific questions",
                "summarize": "Summarize - Get document summaries",
                "compare": "Compare - Compare documents/clauses",
                "risk_analysis": "Risk Analysis - Identify risks"
            }[x]
        )
        
        # Query input
        query = st.text_area("Enter your question:", height=100)
        
        if st.button("Ask Question", type="primary"):
            if query and summary["total_documents"] > 0:
                with st.spinner("Searching documents and generating response..."):
                    result = st.session_state.rag_system.query(query, task_type)
                
                # Add to chat history
                st.session_state.chat_history.append({
                    "question": query,
                    "response": result["response"],
                    "sources": result["sources"],
                    "task_type": task_type,
                    "timestamp": datetime.now().isoformat()
                })
                
                # Display result
                st.subheader("🤖 Response:")
                st.write(result["response"])
                
                # Display sources
                if result["sources"]:
                    st.subheader("📚 Sources:")
                    for i, source in enumerate(result["sources"], 1):
                        with st.expander(f"Source {i}: {source['filename']} (Score: {source['relevance_score']:.3f})"):
                            st.write(f"**Type:** {source['document_type']}")
                            st.write(f"**Preview:** {source['content_preview']}")
            
            elif not query:
                st.warning("Please enter a question.")
            else:
                st.warning("Please upload and process documents first.")
    
    with col2:
        st.header("📋 Quick Actions")
        
        if st.button("📄 Summarize All Documents"):
            if summary["total_documents"] > 0:
                result = st.session_state.rag_system.query("Provide a comprehensive summary of all uploaded documents", "summarize")
                st.write(result["response"])
        
        if st.button("⚠️ Risk Analysis"):
            if summary["total_documents"] > 0:
                result = st.session_state.rag_system.query("Analyze all documents for potential risks", "risk_analysis")
                st.write(result["response"])
        
        if st.button("🔍 Find Key Clauses"):
            if summary["total_documents"] > 0:
                result = st.session_state.rag_system.query("Identify the most important clauses across all documents", "qa")
                st.write(result["response"])
    
    # Chat history
    if st.session_state.chat_history:
        st.header("💬 Chat History")
        for i, chat in enumerate(reversed(st.session_state.chat_history[-5:])):  # Show last 5
            with st.expander(f"Q: {chat['question'][:50]}..." if len(chat['question']) > 50 else f"Q: {chat['question']}"):
                st.write(f"**Task:** {chat['task_type']}")
                st.write(f"**Question:** {chat['question']}")
                st.write(f"**Response:** {chat['response']}")
                st.write(f"**Time:** {chat['timestamp']}")

# FastAPI Backend (Optional)
app = FastAPI(title="Legal Document Q&A API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic models for API
class QueryRequest(BaseModel):
    question: str
    task_type: str = "qa"

class QueryResponse(BaseModel):
    response: str
    sources: List[Dict[str, Any]]
    confidence: float
    task_type: str

# Initialize RAG system
rag_system = LegalRAGSystem()

@app.post("/upload")
async def upload_documents(files: List[UploadFile] = File(...)):
    """Upload and process documents via API"""
    try:
        file_data = []
        for file in files:
            content = await file.read()
            file_data.append((content, file.filename))
        
        results = rag_system.upload_documents(file_data)
        return results
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/query", response_model=QueryResponse)
async def query_documents(request: QueryRequest):
    """Query documents via API"""
    try:
        result = rag_system.query(request.question, request.task_type)
        return QueryResponse(**result)
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/documents")
async def get_documents():
    """Get document summary via API"""
    return rag_system.get_document_summary()

# Main execution
if __name__ == "__main__":
    # Check if running as Streamlit app or FastAPI
    import sys
    
    if "streamlit" in sys.modules:
        # Running as Streamlit app
        create_streamlit_app()
    else:
        # Running as FastAPI
        import uvicorn
        uvicorn.run(app, host="0.0.0.0", port=8000)

